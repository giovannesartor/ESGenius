"""Report generation service — creates PDF and DOCX ESG reports."""

import json
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import UUID

from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.domain.models.report import Report, ReportStatus
from app.repositories.report_repository import ReportRepository
from app.repositories.data_point_repository import DataPointRepository
from app.repositories.company_repository import CompanyRepository
from app.repositories.user_repository import UserRepository
from app.services.scoring_service import ScoringService
from app.services.email_service import EmailService
from app.ai.llm_router import get_router

logger = logging.getLogger(__name__)

_REPORT_SYSTEM_PROMPT = """You are an ESG report writer. Generate a comprehensive ESG report.

Return a JSON object with this structure:
{
  "executive_summary": "string",
  "environmental_section": {
    "overview": "string",
    "key_metrics": [{"metric": "string", "value": "string", "trend": "string"}],
    "highlights": ["string"],
    "risks": ["string"]
  },
  "social_section": {
    "overview": "string",
    "key_metrics": [{"metric": "string", "value": "string", "trend": "string"}],
    "highlights": ["string"],
    "risks": ["string"]
  },
  "governance_section": {
    "overview": "string",
    "key_metrics": [{"metric": "string", "value": "string", "trend": "string"}],
    "highlights": ["string"],
    "risks": ["string"]
  },
  "recommendations": ["string"],
  "methodology_notes": "string"
}"""


def _reports_dir(company_id: UUID) -> Path:
    """Return (and create) the output directory for a company's reports."""
    base = Path(settings.UPLOAD_DIR)
    if not base.is_absolute():
        base = Path(__file__).parent.parent.parent / settings.UPLOAD_DIR
    path = base / "reports" / str(company_id)
    path.mkdir(parents=True, exist_ok=True)
    return path


class ReportService:
    """Generates ESG reports in PDF/DOCX format."""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.report_repo = ReportRepository(db)
        self.dp_repo = DataPointRepository(db)
        self.company_repo = CompanyRepository(db)
        self.scoring = ScoringService(db)
        self.router = get_router()

    async def generate_report(
        self,
        company_id: UUID,
        title: str,
        report_type: str,
        format: str,
        year: int,
        period: str | None,
        generated_by: UUID,
    ) -> Report:
        """Create a report record and trigger generation."""
        report = Report(
            company_id=company_id,
            generated_by=generated_by,
            title=title,
            report_type=report_type,
            format=format,
            year=year,
            period=period,
            status=ReportStatus.GENERATING.value,
        )
        report = await self.report_repo.create(report)
        return report

    async def process_report(self, report_id: UUID) -> None:
        """Process and generate the actual report content."""
        report = await self.report_repo.get_by_id(report_id)
        if not report:
            return

        try:
            company = await self.company_repo.get_by_id(report.company_id)
            data_points = await self.dp_repo.list_by_company(
                report.company_id, year=report.year, limit=10000
            )
            scores = await self.scoring.calculate_scores(
                report.company_id, report.year
            )

            company_info = {
                "name": company.name if company else "Unknown",
                "sector": company.sector if company else None,
                "country": company.country if company else None,
            }
            dp_dicts = [
                {
                    "pillar": dp.pillar,
                    "category": dp.category,
                    "value": dp.value,
                    "numeric_value": dp.numeric_value,
                    "unit": dp.unit,
                    "status": dp.status,
                }
                for dp in data_points
            ]

            user_prompt = (
                f"Company: {json.dumps(company_info)}\n\n"
                f"ESG Data Points: {json.dumps(dp_dicts, indent=2)}\n\n"
                f"ESG Scores: {json.dumps(scores)}"
            )

            # Use the multi-provider router (respects fallback chain + PII redaction)
            ai_result = await self.router.chat_json(
                system=_REPORT_SYSTEM_PROMPT,
                user=user_prompt,
                temperature=0.3,
                max_tokens=4096,
            )

            content_data: dict[str, Any] = ai_result.data or {}
            content_data["_company_name"] = company.name if company else "Company"

            report.content_json = content_data
            report.esg_scores = scores

            if report.format == "pdf":
                file_path = await self._generate_pdf(report, content_data)
            elif report.format == "docx":
                file_path = await self._generate_docx(report, content_data)
            else:
                file_path = None

            report.file_path = file_path
            report.status = ReportStatus.COMPLETED.value
            report.completed_at = datetime.now(timezone.utc)

            try:
                user_repo = UserRepository(self.db)
                user = await user_repo.get_by_id(report.generated_by)
                if user and company:
                    email_service = EmailService()
                    await email_service.send_report_completed_email(
                        email=user.email,
                        name=user.full_name or user.email,
                        report_title=report.title,
                        company_name=company.name,
                        language=getattr(user, "preferred_language", "en"),
                    )
            except Exception as email_err:
                logger.warning("Failed to send report completion email: %s", email_err)

        except Exception as exc:
            logger.error("Report generation failed: %s", exc)
            report.status = ReportStatus.FAILED.value
            report.error_message = str(exc)

        await self.report_repo.update(report)

    # ── PDF generation ────────────────────────────────────────────────────

    async def _generate_pdf(self, report: Report, content: dict) -> str:
        """Generate a professional PDF report using ReportLab."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, PageBreak, KeepTogether,
        )
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

        # ── Colors ──────────────────────────────────────────────────────────
        ESG_GREEN  = colors.HexColor("#16a34a")
        ESG_BLUE   = colors.HexColor("#2563eb")
        ESG_AMBER  = colors.HexColor("#d97706")
        ESG_DARK   = colors.HexColor("#111827")
        ESG_GRAY   = colors.HexColor("#6b7280")
        ESG_LIGHT  = colors.HexColor("#f9fafb")
        ESG_BORDER = colors.HexColor("#e5e7eb")

        PILLAR_COLORS = {
            "Environmental": ESG_GREEN,
            "Social": ESG_BLUE,
            "Governance": ESG_AMBER,
        }
        SECTION_COLORS = [ESG_GREEN, ESG_BLUE, ESG_AMBER]

        output_path = str(_reports_dir(report.company_id) / f"{report.id}.pdf")
        company_name_str = content.get("_company_name", "Company")
        report_year = str(report.year) if report.year else ""

        def on_first_page(canvas, doc):
            canvas.saveState()
            canvas.setFillColor(ESG_GREEN)
            canvas.rect(0, A4[1] - 2 * cm, A4[0], 2 * cm, fill=1, stroke=0)
            canvas.setFillColor(colors.white)
            canvas.setFont("Helvetica-Bold", 16)
            canvas.drawString(2 * cm, A4[1] - 1.35 * cm, "ESG360")
            canvas.setFont("Helvetica", 9)
            canvas.drawRightString(A4[0] - 2 * cm, A4[1] - 1.35 * cm, f"ESG Report · {report_year}")
            canvas.restoreState()

        def on_later_pages(canvas, doc):
            canvas.saveState()
            w, h = A4
            canvas.setFillColor(ESG_GREEN)
            canvas.rect(0, h - 0.7 * cm, w, 0.7 * cm, fill=1, stroke=0)
            canvas.setFillColor(colors.white)
            canvas.setFont("Helvetica-Bold", 8)
            canvas.drawString(2 * cm, h - 0.48 * cm, "ESG360")
            canvas.setFillColor(ESG_LIGHT)
            canvas.rect(0, 0, w, 1.2 * cm, fill=1, stroke=0)
            canvas.setStrokeColor(ESG_BORDER)
            canvas.setLineWidth(0.5)
            canvas.line(0, 1.2 * cm, w, 1.2 * cm)
            canvas.setFillColor(ESG_GRAY)
            canvas.setFont("Helvetica", 7.5)
            canvas.drawString(2 * cm, 0.45 * cm, f"{company_name_str} · ESG Report {report_year} · CONFIDENTIAL")
            canvas.drawRightString(w - 2 * cm, 0.45 * cm, f"Page {doc.page}")
            canvas.restoreState()

        doc = SimpleDocTemplate(
            output_path, pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm,
            topMargin=3 * cm, bottomMargin=2.5 * cm,
            title=report.title, author="ESG360", subject="ESG Report",
        )

        base = getSampleStyleSheet()

        def S(name, **kw):
            return ParagraphStyle(name, **kw)

        sTitle   = S("RTitle",  parent=base["Title"],   fontSize=36, textColor=ESG_DARK, spaceAfter=8,  leading=44, alignment=TA_CENTER, fontName="Helvetica-Bold")
        sSubtitle= S("RSub",    parent=base["Normal"],  fontSize=14, textColor=ESG_GRAY, spaceAfter=4,  leading=20, alignment=TA_CENTER)
        sH1      = S("RH1",     parent=base["Heading1"],fontSize=18, textColor=ESG_DARK, spaceBefore=18,spaceAfter=8,  fontName="Helvetica-Bold")
        sH2      = S("RH2",     parent=base["Heading2"],fontSize=13, textColor=ESG_DARK, spaceBefore=12,spaceAfter=6,  fontName="Helvetica-Bold")
        sBody    = S("RBody",   parent=base["Normal"],  fontSize=10, textColor=ESG_DARK, spaceAfter=6,  leading=15, alignment=TA_JUSTIFY)
        sCaption = S("RCap",    parent=base["Normal"],  fontSize=8,  textColor=ESG_GRAY, spaceAfter=4,  leading=12)
        sBullet  = S("RBullet", parent=base["Normal"],  fontSize=10, textColor=ESG_DARK, spaceAfter=4,  leading=14, leftIndent=14, firstLineIndent=-10)
        sLabel   = S("RLabel",  parent=base["Normal"],  fontSize=8,  textColor=ESG_GRAY, spaceAfter=2,  fontName="Helvetica-Bold")
        sValue   = S("RValue",  parent=base["Normal"],  fontSize=22, textColor=ESG_GREEN,spaceAfter=2,  fontName="Helvetica-Bold", alignment=TA_CENTER)
        sMeta    = S("RMeta",   parent=base["Normal"],  fontSize=9,  textColor=ESG_GRAY, spaceAfter=4,  alignment=TA_CENTER)

        story = []

        # ── Cover ────────────────────────────────────────────────────────────
        story.append(Spacer(1, 3 * cm))
        story.append(HRFlowable(width="100%", thickness=6, color=ESG_GREEN, spaceAfter=28))
        story.append(Paragraph(report.title or "ESG Report", sTitle))
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph(company_name_str, sSubtitle))
        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph(f"Reporting Year: {report_year}  ·  Framework: {report.report_type or 'Multi-Framework'}", sMeta))
        story.append(Spacer(1, 1 * cm))

        scores = report.esg_scores or {}
        if scores:
            score_data = [
                [Paragraph("Overall", sLabel), Paragraph("Environmental", sLabel), Paragraph("Social", sLabel), Paragraph("Governance", sLabel)],
                [
                    Paragraph(str(int(scores.get("overall", 0))), sValue),
                    Paragraph(str(int(scores.get("environmental", 0))), S("RVE", parent=sValue, textColor=ESG_GREEN)),
                    Paragraph(str(int(scores.get("social", 0))),        S("RVS", parent=sValue, textColor=ESG_BLUE)),
                    Paragraph(str(int(scores.get("governance", 0))),    S("RVG", parent=sValue, textColor=ESG_AMBER)),
                ],
            ]
            score_table = Table(score_data, colWidths=["25%", "25%", "25%", "25%"])
            score_table.setStyle(TableStyle([
                ("BACKGROUND",  (0, 0), (-1, -1), ESG_LIGHT),
                ("BACKGROUND",  (0, 0), (-1,  0), colors.white),
                ("BOX",         (0, 0), (-1, -1), 1, ESG_BORDER),
                ("INNERGRID",   (0, 0), (-1, -1), 0.5, ESG_BORDER),
                ("TOPPADDING",  (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING",(0,0), (-1, -1), 10),
                ("ALIGN",       (0, 0), (-1, -1), "CENTER"),
                ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(score_table)
            story.append(Spacer(1, 0.5 * cm))

        story.append(HRFlowable(width="100%", thickness=1, color=ESG_BORDER, spaceAfter=12))
        from datetime import date
        story.append(Paragraph(f"Generated by ESG360 AI Platform  ·  {date.today().strftime('%B %d, %Y')}  ·  CONFIDENTIAL", sMeta))
        story.append(PageBreak())

        # ── TOC ──────────────────────────────────────────────────────────────
        story.append(Paragraph("Table of Contents", sH1))
        story.append(HRFlowable(width="100%", thickness=1.5, color=ESG_GREEN, spaceAfter=12))
        toc_items = [
            ("1.", "Executive Summary"),
            ("2.", "Environmental Performance"),
            ("3.", "Social Performance"),
            ("4.", "Governance Performance"),
            ("5.", "Recommendations"),
            ("6.", "Methodology & Sources"),
        ]
        for num, title_text in toc_items:
            toc_row = Table(
                [[Paragraph(f"<b>{num}</b>", sBody), Paragraph(title_text, sBody), Paragraph("", sCaption)]],
                colWidths=[1 * cm, 12 * cm, None],
            )
            toc_row.setStyle(TableStyle([
                ("ALIGN",  (2, 0), (2, 0), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING",    (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(toc_row)
        story.append(PageBreak())

        # ── Executive Summary ────────────────────────────────────────────────
        story.append(Paragraph("1. Executive Summary", sH1))
        story.append(HRFlowable(width="100%", thickness=1.5, color=ESG_GREEN, spaceAfter=12))
        if content.get("executive_summary"):
            story.append(Paragraph(content["executive_summary"], sBody))
        story.append(Spacer(1, 0.5 * cm))

        # ── E / S / G Sections ───────────────────────────────────────────────
        section_map = [
            ("environmental_section", "2. Environmental Performance", SECTION_COLORS[0]),
            ("social_section",        "3. Social Performance",        SECTION_COLORS[1]),
            ("governance_section",    "4. Governance Performance",    SECTION_COLORS[2]),
        ]
        score_keys = ["environmental", "social", "governance"]

        for (section_key, section_title, section_color), score_key in zip(section_map, score_keys):
            section = content.get(section_key, {})
            if not section:
                continue

            story.append(Paragraph(section_title, sH1))
            story.append(HRFlowable(width="100%", thickness=1.5, color=section_color, spaceAfter=12))

            pillar_score = int(scores.get(score_key, 0))
            if pillar_score:
                badge = Table(
                    [[Paragraph(f"Score: {pillar_score}/100", S("badge", parent=sBody, textColor=colors.white, fontName="Helvetica-Bold"))]],
                    colWidths=[4 * cm],
                )
                badge.setStyle(TableStyle([
                    ("BACKGROUND",    (0, 0), (-1, -1), section_color),
                    ("TOPPADDING",    (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("LEFTPADDING",   (0, 0), (-1, -1), 12),
                ]))
                story.append(badge)
                story.append(Spacer(1, 0.3 * cm))

            if section.get("overview"):
                story.append(Paragraph(section["overview"], sBody))

            metrics = section.get("key_metrics", [])
            if metrics:
                story.append(Spacer(1, 0.3 * cm))
                story.append(Paragraph("Key Metrics", sH2))
                header = [Paragraph("<b>Metric</b>", sCaption), Paragraph("<b>Value</b>", sCaption), Paragraph("<b>Trend</b>", sCaption)]
                rows = [header] + [
                    [Paragraph(m.get("metric", ""), sBody), Paragraph(str(m.get("value", "")), sBody), Paragraph(m.get("trend", "—"), sBody)]
                    for m in metrics[:8]
                ]
                metrics_table = Table(rows, colWidths=[8 * cm, 4 * cm, 3 * cm])
                metrics_table.setStyle(TableStyle([
                    ("BACKGROUND",    (0, 0), (-1,  0), ESG_LIGHT),
                    ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, ESG_LIGHT]),
                    ("BOX",           (0, 0), (-1, -1), 0.5, ESG_BORDER),
                    ("INNERGRID",     (0, 0), (-1, -1), 0.25, ESG_BORDER),
                    ("TOPPADDING",    (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                    ("LEFTPADDING",   (0, 0), (-1, -1), 8),
                    ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ]))
                story.append(metrics_table)

            for label, items_key in [("Highlights", "highlights"), ("Risks & Challenges", "risks")]:
                items = section.get(items_key, [])
                if items:
                    story.append(Spacer(1, 0.3 * cm))
                    story.append(Paragraph(label, sH2))
                    prefix = "•" if items_key == "highlights" else "▸"
                    for item in items:
                        story.append(Paragraph(f"{prefix} {item}", sBullet))

            story.append(PageBreak())

        # ── Recommendations ──────────────────────────────────────────────────
        recs = content.get("recommendations", [])
        if recs:
            story.append(Paragraph("5. Recommendations", sH1))
            story.append(HRFlowable(width="100%", thickness=1.5, color=ESG_GREEN, spaceAfter=12))
            for i, rec in enumerate(recs, 1):
                # Use ESG_GREEN consistently — no fragile variable lookup
                rec_box = Table(
                    [[Paragraph(f"<b>{i:02d}</b>", S("rnum", parent=sValue, fontSize=14, textColor=ESG_GREEN)), Paragraph(rec, sBody)]],
                    colWidths=[1.2 * cm, None],
                )
                rec_box.setStyle(TableStyle([
                    ("BACKGROUND",    (0, 0), (-1, -1), ESG_LIGHT),
                    ("TOPPADDING",    (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING",   (0, 0), (-1, -1), 12),
                    ("RIGHTPADDING",  (0, 0), (-1, -1), 12),
                    ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                    ("BOX",           (0, 0), (-1, -1), 0.5, ESG_BORDER),
                ]))
                story.append(rec_box)
                story.append(Spacer(1, 0.3 * cm))
            story.append(PageBreak())

        # ── Methodology ──────────────────────────────────────────────────────
        story.append(Paragraph("6. Methodology & Sources", sH1))
        story.append(HRFlowable(width="100%", thickness=1.5, color=ESG_GREEN, spaceAfter=12))
        methodology = content.get("methodology_notes") or (
            "This report was generated by the ESG360 AI platform. Data was extracted from company-provided documents "
            "using advanced Natural Language Processing (NLP) and classified according to the GRI Standards, SASB "
            "Frameworks, and TCFD Recommendations. ESG scores are calculated using a weighted methodology across "
            "Environmental (40%), Social (30%), and Governance (30%) pillars, with sub-indicator scoring based on "
            "data availability, validation status, and AI confidence levels. Industry benchmarks are derived from "
            "publicly available ESG disclosure data and representative sector baselines."
        )
        story.append(Paragraph(methodology, sBody))
        story.append(Spacer(1, 0.5 * cm))

        disclaimer_box = Table(
            [[Paragraph(
                "<b>Disclaimer:</b> This report is based on data provided by the reporting organization. "
                "ESG360 does not independently verify the accuracy of underlying data. This report is intended "
                "for informational purposes and should not be construed as investment, legal, or regulatory advice.",
                S("disc", parent=sCaption, textColor=ESG_GRAY),
            )]],
            colWidths=["100%"],
        )
        disclaimer_box.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, -1), ESG_LIGHT),
            ("BOX",           (0, 0), (-1, -1), 0.5, ESG_BORDER),
            ("TOPPADDING",    (0, 0), (-1, -1), 12),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
            ("LEFTPADDING",   (0, 0), (-1, -1), 14),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 14),
        ]))
        story.append(disclaimer_box)

        doc.build(story, onFirstPage=on_first_page, onLaterPages=on_later_pages)
        return output_path

    # ── DOCX generation ───────────────────────────────────────────────────

    async def _generate_docx(self, report: Report, content: dict) -> str:
        """Generate a styled DOCX report using python-docx."""
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        output_path = str(_reports_dir(report.company_id) / f"{report.id}.docx")

        doc = DocxDocument()

        # ── Page margins ─────────────────────────────────────────────────
        sections = doc.sections
        for section in sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # ── Helpers ──────────────────────────────────────────────────────
        GREEN = RGBColor(0x16, 0xa3, 0x4a)
        DARK  = RGBColor(0x11, 0x18, 0x27)
        GRAY  = RGBColor(0x6b, 0x72, 0x80)

        def add_heading(text: str, level: int = 1, color: RGBColor = DARK) -> None:
            h = doc.add_heading(text, level=level)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                run.font.color.rgb = color
                run.font.bold = True

        def add_paragraph_text(text: str, bold: bool = False, size: int = 10) -> None:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = DARK

        def add_metric_table(metrics: list[dict]) -> None:
            if not metrics:
                return
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            for cell, label in zip(hdr, ["Metric", "Value", "Trend"]):
                cell.text = label
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
            for m in metrics[:8]:
                row = tbl.add_row().cells
                row[0].text = m.get("metric", "")
                row[1].text = str(m.get("value", ""))
                row[2].text = m.get("trend", "—")
                for cell in row:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)
            doc.add_paragraph()

        company_name_str = content.get("_company_name", "Company")
        report_year = str(report.year) if report.year else ""

        # ── Cover ────────────────────────────────────────────────────────
        cover_title = doc.add_heading(report.title or "ESG Report", level=0)
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cover_title.runs:
            run.font.color.rgb = GREEN
            run.font.bold = True

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run(f"{company_name_str}  ·  {report_year}")
        r.font.size = Pt(13)
        r.font.color.rgb = GRAY

        # ESG Scores summary
        scores = report.esg_scores or {}
        if scores:
            doc.add_paragraph()
            tbl = doc.add_table(rows=2, cols=4)
            tbl.style = "Table Grid"
            headers = ["Overall", "Environmental", "Social", "Governance"]
            vals = [
                str(int(scores.get("overall", 0))),
                str(int(scores.get("environmental", 0))),
                str(int(scores.get("social", 0))),
                str(int(scores.get("governance", 0))),
            ]
            for i, (hdr_text, val_text) in enumerate(zip(headers, vals)):
                tbl.rows[0].cells[i].text = hdr_text
                for run in tbl.rows[0].cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                tbl.rows[1].cells[i].text = val_text
                for run in tbl.rows[1].cells[i].paragraphs[0].runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.color.rgb = GREEN

        doc.add_page_break()

        # ── Executive Summary ────────────────────────────────────────────
        add_heading("1. Executive Summary", level=1, color=GREEN)
        if content.get("executive_summary"):
            add_paragraph_text(content["executive_summary"])
        doc.add_paragraph()

        # ── E / S / G Sections ───────────────────────────────────────────
        section_map = [
            ("environmental_section", "2. Environmental Performance"),
            ("social_section",        "3. Social Performance"),
            ("governance_section",    "4. Governance Performance"),
        ]
        score_keys = ["environmental", "social", "governance"]

        for (section_key, section_title), score_key in zip(section_map, score_keys):
            section = content.get(section_key, {})
            if not section:
                continue

            add_heading(section_title, level=1, color=GREEN)
            pillar_score = int(scores.get(score_key, 0))
            if pillar_score:
                p = doc.add_paragraph()
                r = p.add_run(f"Score: {pillar_score}/100")
                r.font.bold = True
                r.font.color.rgb = GREEN

            if section.get("overview"):
                add_paragraph_text(section["overview"])

            metrics = section.get("key_metrics", [])
            if metrics:
                add_heading("Key Metrics", level=2)
                add_metric_table(metrics)

            for label, key in [("Highlights", "highlights"), ("Risks & Challenges", "risks")]:
                items = section.get(key, [])
                if items:
                    add_heading(label, level=2)
                    for item in items:
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(item).font.size = Pt(10)

            doc.add_page_break()

        # ── Recommendations ──────────────────────────────────────────────
        recs = content.get("recommendations", [])
        if recs:
            add_heading("5. Recommendations", level=1, color=GREEN)
            for i, rec in enumerate(recs, 1):
                p = doc.add_paragraph(style="List Number")
                p.add_run(rec).font.size = Pt(10)
            doc.add_page_break()

        # ── Methodology ──────────────────────────────────────────────────
        add_heading("6. Methodology & Sources", level=1, color=GREEN)
        methodology = content.get("methodology_notes") or (
            "This report was generated by the ESG360 AI platform using GRI Standards, "
            "SASB Frameworks, and TCFD Recommendations as reference frameworks."
        )
        add_paragraph_text(methodology)

        disclaimer = doc.add_paragraph()
        r = disclaimer.add_run(
            "Disclaimer: This report is based on data provided by the reporting organization. "
            "ESG360 does not independently verify the accuracy of underlying data."
        )
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        r.italic = True

        doc.save(output_path)
        return output_path
